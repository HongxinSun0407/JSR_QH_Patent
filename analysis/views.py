import io
import json
import os
import re
import shutil
import string
import subprocess
import time
import uuid
import zipfile
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from itertools import zip_longest
import openpyxl
import pandas as pd
import requests
from celery import shared_task
from celery.bin.control import status
from dateutil.relativedelta import relativedelta
from django.conf import settings
from django.core.files import File
from django.db.models import Q
from django.db import transaction
from django.db.models.functions import Cast
from django.http import FileResponse, HttpResponse
from docx import Document
from docx.enum.text import WD_BREAK
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import Inches
from openpyxl.styles import Alignment
from pymupdf import pymupdf
from rest_framework import generics, viewsets, mixins
from rest_framework.decorators import action
from rest_framework.exceptions import NotFound, PermissionDenied
from rest_framework.pagination import PageNumberPagination
from rest_framework.permissions import IsAdminUser
from rest_framework.response import Response
from rest_framework.views import APIView

from analysis.export_excel import export_excel
from analysis.serializers import *
from analysis.tiangong_ai import request_tiangong_chat_ai, request_tiangong_copilot_ai, request_tiangong_research_ai, \
    request_tiangong_image_ai, get_tokens_to_model
from file.models import FileContentModel, FileModel, FileSerializer
from file.views import get_file_prefix, upload_file_service
from patent_ai.exceptions import logger
from patent_ai.settings.base import KIMI_TIMEOUT, TEMP_KEY_NAME_LIST, PROBLEM_MAPPING_KEY_JSON, \
    VERSION_PREFIX, df, class_name_df
from user.permissions import CustomModelPermissions
from user.serializers import UserSerializer

re_str = r"[-\s\*]+[省份城市地址点总部位置所]{2,6}\**[:：]"

class ZipViewSet(generics.ListCreateAPIView, mixins.RetrieveModelMixin,viewsets.GenericViewSet):
    permission_classes = [CustomModelPermissions]
    serializer_class = ZipAnalysisSerializer

    def get_queryset(self):
        user = self.request.user
        if user.is_superuser:
            return ZipAnalysisModel.objects.all()
        else:
            return ZipAnalysisModel.objects.filter(user_id=user)


class ChatContentResult(APIView):
    permission_classes = [IsAdminUser]

    def get(self, request, session_id):
        serializer = ChatContentSerializer(ChatContentModel.objects.filter(session_id=session_id),
                                           many=True)
        return Response(serializer.data)

class ScoreView(APIView):
    permission_classes = [IsAdminUser]
    def get(self, request):

        zip_result_list = ZipAnalysisResultModel.objects.filter(status='success', patent_info__isnull=False)
        # 加载现有的 Excel 文件
        workbook = openpyxl.load_workbook('评分模型.xlsx')

        # 选择一个工作表
        market = workbook['市场维度']
        tech = workbook['技术维度']
        law = workbook['法律维度']
        total = workbook['总分']
        start_index = 3
        for index, zip_result in enumerate(zip_result_list):
            score = ChatContentModel.objects.filter(problem_label_id=12, role='system',session_id=zip_result.session_id.id).values("content")[0]
            industry = ChatContentModel.objects.filter(problem_label_id=16, role='system',
                                            session_id=zip_result.session_id.id).values("content")[0]['content']
            patent_json = zip_result.patent_info
            score_json = score['content']
            p_code = patent_json.get("申请公布号", "").replace(" ", "")[:-1]
            score_json['申请公布号'] = patent_json.get("申请公布号", "")
            compute_score(score_json, zip_result)
            patent_info = df[df['公开(公告)号'] == p_code + "B"]
            if patent_info.empty:
                patent_info = df[df['公开(公告)号'] == p_code + "A"]
            if patent_info.empty:
                patent_info = df[df['公开(公告)号'] == p_code + "C"]
            if patent_info.empty:
                patent_info = df[df['公开(公告)号'] == p_code + "U"]
            if patent_info.empty:
                patent_info = df[df['公开(公告)号'] == p_code + "S"]
            filtered_data1 = patent_info['一级产业分类']
            filtered_data2 = patent_info['二级产业分类']
            industry_str = industry.get("行业分类")
            # pattern1 = re.compile(r'.*?【战略新兴产业分类1】[:：](.*?)[;；,，\s]{1}(.*)')
            pattern2 = re.compile(r'.*?【战略新兴产业分类2】[:：](.*?)[;；,，\s]{1}(.*)')
            # match1 = pattern1.search(industry_str)
            match2 = pattern2.search(industry_str)
            type1 = filtered_data1.iloc[0] if not filtered_data1.empty else '无'
            type1_2 = filtered_data2.iloc[0] if not filtered_data2.empty else '无'
            type2 = ""
            type2_2 = ""
            # if match1:
            #     # 提取分组的值
            #     type1 = match1.group(1)
            #     type1_2 = match1.group(2)
            if match2:
                type2 = match2.group(1)
                type2_2 = match2.group(2)
            total['A' + str(start_index-1 + index)] = str(index+1)
            total['B' + str(start_index-1 + index)] = patent_json['专利名称']
            total['C' + str(start_index-1 + index)] = patent_json['专利申请号']
            total['D' + str(start_index - 1 + index)] = type1
            total['E' + str(start_index - 1 + index)] = type1_2
            total['F' + str(start_index - 1 + index)] = type2
            total['G' + str(start_index - 1 + index)] = type2_2
            total['H' + str(start_index-1 + index)] = f"=ROUND(I{start_index-1 + index}*评分模型!$B$2+J{start_index-1 + index}*评分模型!$B$10+K{start_index-1 + index}*评分模型!$B$19,1)"
            total['I' + str(start_index-1 + index)] = f"=ROUND(法律维度!D{start_index + index}*评分模型!$E$2+法律维度!I{start_index + index}*评分模型!$E$6,1)"
            total['J' + str(start_index-1 + index)] = f"=ROUND(技术维度!D{start_index + index}*评分模型!$E$10+技术维度!H{start_index + index}*评分模型!$E$13+技术维度!M{start_index + index}*评分模型!$E$17,1)"
            total['K' + str(start_index-1 + index)] = f"=ROUND(市场维度!D{start_index + index}*评分模型!$E$19+市场维度!G{start_index + index}*评分模型!$E$21+市场维度!L{start_index + index}*评分模型!$E$25,1)"
            clazz_type_data = get_clazz_type(p_code)
            code_list = [ "L", "M",
                    "N", "O", "P", "Q", "R", "S", "T", "U", "V", "W", "X", "Y", "Z",
                    "AA", "AB", "AC", "AD", "AE", "AF", "AG", "AH", "AI", "AJ", "AK",
                    "AL", "AM", "AN", "AO", "AP", "AQ", "AR", "AS"
                ]
            count = 1
            code_index = 0
            for clazz_index,clazz_value in enumerate(clazz_type_data.values()):
                code_key = code_list[code_index]
                total[code_key + str(start_index - 2)] = f"国民经济行业分类{count}门类"
                total[code_key + str(start_index-1 + index)] = clazz_value['big_name']
                code_index += 1
                code_key = code_list[code_index]
                total[code_key + str(start_index - 2)] = f"国民经济行业分类{count}大类"
                category_str = "；".join(clazz_value['category_list'])
                total[code_key + str(start_index-1 + index)] = category_str
                code_index+=1
                count += 1


            if score_json:
                market['A' + str(start_index + index)] = str(index + 1)
                market['B' + str(start_index + index)] = patent_json['专利名称']
                market['C' + str(start_index + index)] = patent_json['专利申请号']
                market['D' + str(start_index + index)] = f"=ROUND((E{start_index+index}*评分模型!$H$19)+(F{start_index+index}*评分模型!$H$20),1)"
                market['E' + str(start_index+index)] = score_json['潜在市场规模']
                market['F' + str(start_index + index)] = score_json['竞争格局']
                market['G' + str(
                    start_index + index)] = f"=ROUND(H{start_index+index}*评分模型!$H$21+I{start_index+index}*评分模型!$H$22+J{start_index+index}*评分模型!$H$23+K{start_index+index}*评分模型!$H$24,1)"
                market['H' + str(start_index + index)] = score_json['预期利润']
                market['I' + str(start_index + index)] = score_json['市场份额']
                market['J' + str(start_index + index)] = score_json['年限分数']
                market['K' + str(start_index + index)] = score_json['当前申请人分数']
                market['L' + str(
                    start_index + index)] = f"=ROUND(M{start_index+index}*评分模型!$H$25+N{start_index+index}*评分模型!$H$26,1)"
                market['M' + str(start_index + index)] = score_json['渠道拓展与维护难度']
                market['N' + str(start_index + index)] = score_json['营销资源投入']

                tech['A' + str(start_index + index)] = str(index + 1)
                tech['B' + str(start_index + index)] = patent_json['专利名称']
                tech['C' + str(start_index + index)] = patent_json['专利申请号']
                tech['D' + str(start_index + index)] = f"=ROUND(E{start_index+index}*评分模型!$H$10+F{start_index+index}*评分模型!$H$11+G{start_index+index}*评分模型!$H$12,1)"
                tech['E' + str(start_index + index)] = score_json['技术创新性']
                tech['F' + str(start_index + index)] = score_json['技术独特性']
                tech['G' + str(start_index + index)] = score_json['被引用分数']

                tech['H' + str(start_index + index)] = f"=ROUND(I{start_index+index}*评分模型!$H$13+J{start_index+index}*评分模型!$H$14+K{start_index+index}*评分模型!$H$15+L{start_index+index}*评分模型!$H$16,1)"
                tech['I' + str(start_index + index)] = score_json['技术复杂性']
                tech['J' + str(start_index + index)] = score_json['资源需求']
                tech['K' + str(start_index + index)] = score_json['技术转化周期']
                tech['L' + str(start_index + index)] = score_json['风险与不确定性']
                tech['M' + str(start_index + index)] = f"=ROUND(N{start_index+index}*评分模型!$H$17+O{start_index+index}*评分模型!$H$18,1)"
                tech['N' + str(start_index + index)] = score_json['社会与经济影响']
                tech['O' + str(start_index + index)] = score_json['战略新兴产业分类得分']


                law['A' + str(start_index + index)] = str(index + 1)
                law['B' + str(start_index + index)] = patent_json['专利名称']
                law['C' + str(start_index + index)] = patent_json['专利申请号']
                law['D' + str(start_index + index)] = f"=ROUND((E{start_index+index}*评分模型!$H$2)+(F{start_index+index}*评分模型!$H$3)+(G{start_index+index}*评分模型!$H$4)+(H{start_index+index}*评分模型!$H$5),1)"
                law['E' + str(start_index + index)] = score_json['权利要求的清晰性']
                law['F' + str(start_index + index)] = score_json.get("保护范围的深度",score_json.get("专利对核心技术细节的保护程度"))
                law['G' + str(start_index + index)] = score_json['说明书分数']
                law['H' + str(start_index + index)] = score_json['独立权利要求分数']

                law['I' + str(start_index + index)] = f"=ROUND(J{start_index+index}*评分模型!$H$6+K{start_index+index}*评分模型!$H$7+L{start_index+index}*评分模型!$H$8+M{start_index+index}*评分模型!$H$9,1)"
                law['J' + str(start_index + index)] = score_json['受到技术挑战的可能性']
                law['K' + str(start_index + index)] = score_json['专利类型得分']
                law['L' + str(start_index + index)] = score_json['简单同族分数']
                law['M' + str(start_index + index)] = score_json['法律状态得分']

        # 创建一个对齐对象，用于居中对齐
        alignment = Alignment(horizontal='center', vertical='center')

        # 遍历工作簿中的所有工作表
        for sheet in workbook.worksheets:
            # 获取工作表的最大行数和列数
            max_row = sheet.max_row
            max_col = sheet.max_column

            # 对当前工作表的所有单元格进行居中对齐
            for row in sheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
                for cell in row:
                    cell.alignment = alignment
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=workbook.xlsx'
        excel_stream = io.BytesIO()
        workbook.save(excel_stream)
        excel_stream.seek(0)
        # 将工作簿保存到HttpResponse中
        response.write(excel_stream.getvalue())
        return response

    def post(self, request, *args, **kwargs):
        zip_result_id = request.data.get("zip_result_id")
        zarm = ZipAnalysisResultModel.objects.get(pk=zip_result_id)
        data = zarm.patent_info
        result = zarm.session_id.chatcontentmodel_set.filter(problem_label_id=12, role='system').first().content
        result['申请公布号'] = data['申请公布号']
        compute_score(result, zarm)
        result.pop("申请公布号")
        return Response(result)


class ProblemLabelList(mixins.ListModelMixin,
                       mixins.UpdateModelMixin,
                       viewsets.GenericViewSet):
    queryset = ProblemLabelModel.objects.all()
    permission_classes = [CustomModelPermissions]
    serializer_class = ProblemLabelSerializer

    def get_queryset(self):
        user = self.request.user
        if user.is_superuser:
            return ProblemLabelModel.objects.all()
        else:
            return ProblemLabelModel.objects.all().filter(id=-1)


class ResultExportView(
    mixins.RetrieveModelMixin,
    mixins.UpdateModelMixin,
    mixins.ListModelMixin,
    viewsets.GenericViewSet):
    queryset = ResultExportModel.objects.all()
    permission_classes = [CustomModelPermissions]
    serializer_class = ResultExportSerializer
    def get_queryset(self):
        user = self.request.user
        username = user.username
        if user.is_superuser:
            zar_query = ZipAnalysisResultModel.objects.all()
        else:
            zar_query = ZipAnalysisResultModel.objects.filter(Q(patent_info__发明人__icontains=username) | Q(user_id=user))
        result = ResultExportModel.objects.filter(zip_analysis_id__in=zar_query).order_by("-update_time").all()
        return result
    def list(self, request, *args, **kwargs):
        result = self.serializer_class(self.get_queryset(), many=True)
        return Response(result.data)
    @action(detail=False, methods=['put'], url_path=ResultExportModel._meta.permissions[0][0])
    @transaction.atomic
    def batch_update(self, request):
        pdated_objects = []
        for item in request.data:
            result_export = self.get_queryset().get(id=item['id'])
            res = self.get_serializer(result_export, data=item)
            if res.is_valid(raise_exception=True):
                res.save()
                pdated_objects.append(res.data)
        return Response(pdated_objects)

    @action(detail=False, methods=['get'], url_path=ResultExportModel._meta.permissions[1][0])
    def export(self,request):
        type = request.query_params.get("id")
        data = self.get_queryset()
        if type == '1':
            data = [ResultExportModel(zip_analysis_id=i) for i in ZipAnalysisResultModel.objects.filter(status='success').all()]
        teacher_name_pd = pd.read_excel('老师名单.xlsx')
        export_list = []
        for i in data:
            try:
                zarm = ZipAnalysisResultModel.objects.get(id=i.zip_analysis_id.id)
                if zarm.patent_info:
                    inventor_list = re.split(r'[,\s、]+', zarm.patent_info.get('发明人', "").strip())
                    department_set = set()
                    for inventor_name in inventor_list:
                        teacher_col = teacher_name_pd[teacher_name_pd['姓名'] == inventor_name.strip()]
                        if not teacher_col.empty:
                            department_set.add(teacher_col['聘任部门'].iloc[0])
                    i.department = ",".join(department_set)
                    ccm:ChatContentModel = ChatContentModel.objects.filter(session_id= zarm.session_id.id, problem_label_id=12,role='system').first()
                    score_map = get_score(zarm.id, ccm.content)
                    i.law_score = score_map['法律得分']
                    i.market_score = score_map['市场得分']
                    i.tech_score = score_map['技术得分']
                    i.ai_score = score_map['综合评分']
                    p_code = zarm.patent_info['申请公布号'].replace(" ", '')[:-1]
                    i.problem_solved = zarm.patent_info['解决问题']
                    patent_info = df[df['公开(公告)号'] == p_code + "B"]
                    if patent_info.empty:
                        patent_info = df[df['公开(公告)号'] == p_code + "A"]
                    if patent_info.empty:
                        patent_info = df[df['公开(公告)号'] == p_code + "C"]
                    if patent_info.empty:
                        patent_info = df[df['公开(公告)号'] == p_code + "U"]
                    if patent_info.empty:
                        patent_info = df[df['公开(公告)号'] == p_code + "S"]
                    if not patent_info.empty:
                        legal_status = patent_info['简单法律状态'].iloc[0]
                        authorization_date = pd.to_datetime(patent_info['授权日'].iloc[0])
                        i.apply_code = patent_info['申请号'].iloc[0]
                        i.inventor = patent_info['发明人'].iloc[0]
                        i.patent_name = patent_info['标题'].iloc[0]
                        i.legal_status = legal_status
                        apply_date = patent_info['申请日'].iloc[0]
                        difference = relativedelta(pd.to_datetime(datetime.now()), pd.to_datetime(apply_date))
                        i.maintenance_period = difference.years
                        type_dict = get_clazz_type(p_code)
                        neic = ""
                        for clazz_index, value in enumerate(type_dict.values()):
                            if len(type_dict) > 1:
                                neic += f"国民经济行业分类{clazz_index + 1}\n"
                            else:
                                neic += "国民经济行业分类\n"
                            category_str = "；".join(value['category_list'])
                            neic += f"门类：{value['big_name']}\n大类：{category_str}\n"
                        i.neic = neic
                        i.new_classification = patent_info['战略新兴产业分类'].iloc[0]
                        i.ipc = patent_info['IPC分类号'].iloc[0]
                        i.neic_num = patent_info['国民经济行业分类号'].iloc[0]
                        i.technical_topic_class = patent_info['技术主题分类'].iloc[0]
                        i.application_field_class = patent_info['应用领域分类'].iloc[0]
                    else:
                        i.apply_code = "CN"+zarm.patent_info['专利申请号']
                        i.inventor = zarm.patent_info['发明人']
                        i.patent_name = zarm.patent_info['专利名称']
                        continue
                    ccm = ChatContentModel.objects.filter(session_id = zarm.session_id.id,problem_label_id=7,role='system').first()
                    if ccm:
                        i.ctp = ccm.content.get("产业化前景","").replace("#","")
                    ccm = ChatContentModel.objects.filter(session_id = zarm.session_id.id,problem_label_id=5,role='system').first()
                    if ccm:
                        doip = ccm.content.get("本发明与现有技术的比较分析")
                        doip = process_object({"":doip})
                        i.doip = doip
                    ccm = ChatContentModel.objects.filter(session_id=zarm.session_id.id, problem_label_id=11,
                                                          role='system').first()
                    if ccm :
                        data_markdown = ccm.content.get("合作企业")
                        en_data = [i.strip() for i in data_markdown.split("\n") if i != '']
                        c = []
                        d = set()
                        for line in en_data:
                            if re.match(r'^\d.', line):
                                line = line.replace("*", "")
                                d_item = re.findall(r'[（\(](.*?)[\)）]', line)
                                if d_item:
                                    d.add(d_item[0])
                                d_item = re.findall(r'-\s*(.*?市)', line)
                                if d_item:
                                    d.add(d_item[0])
                                c.append(line)
                            if re.match(re_str, line):
                                d.add(re.sub(re_str, "", line))
                        en_data = [i for i in zip_longest(c, d, fillvalue='')]
                        i.cooperative_enterprises = ", ".join([i[0] for i in en_data])
                        i.cooperative_city = ",".join([i[1] for i in en_data if i[1] != ''])
                    export_list.append(i)
            except Exception as e:
                logger.error("出现错误", exc_info=True)
        if type == '1':
            file_path = export_excel(export_list, template_excel="盘活行动模板1.xlsx")
        else:
            file_path = export_excel(export_list)
        return FileResponse(open(file_path, 'rb'), as_attachment=False)


    @action(detail=False, methods=['get'], url_path="get_file_name")
    def get_file_name_reuqest(self,request):
        user = request.user
        data = self.get_queryset()
        if len(data) > 0:
            first = data[0]
            patent_name = data[0].patent_name
            count = len(data)
            formatted_number = f"{VERSION_PREFIX}{first.zip_analysis_id.id:04}"
            if len(data) == 1:
                file_name = f'{formatted_number}-{patent_name}-{user.username}'
            else:
                file_name = f'{formatted_number}等{count}件专利-{patent_name}等{count}件专利-{user.username}'
        else:
            raise NotFound(detail="没有数据")
        return Response(file_name)

class StandardResultsSetPagination(PageNumberPagination):
    page_size = 100
    page_size_query_param = 'perPage'
    max_page_size = 1000

class ZipResultViewSet(mixins.RetrieveModelMixin,mixins.ListModelMixin, viewsets.GenericViewSet):
    queryset = ZipAnalysisResultModel.objects.all()
    serializer_class = ZipAnalysisResultSerializer
    permission_classes = [CustomModelPermissions]
    pagination_class = StandardResultsSetPagination
    def get_queryset(self):
        user = self.request.user
        if user.is_superuser:
            zar_query = ZipAnalysisResultModel.objects.all()
        else:
            zar_query = ZipAnalysisResultModel.objects.filter(user_id=user).all()
        return zar_query

    def list(self, request, *args, **kwargs):
        data = request_moonshot_ai({"messages": [{"role": "user",
                                                  "content": "1.提取专利的基础信息。\n2.按照传统行业分类和战略新兴产业分类两种方式，列出专利应用领域所属的分类。\n3.请根据文本提供的信息判断本专利的法律状态，要求为“尚未授权、申请撤回、申请被驳回、实质审查、专利权无效、已授权、专利权终止、专利权无效”当中一种状态。\n4.根据专利文本的技术背景和发明内容最后一段，归纳出发明【解决问题】，要求字数200-300字之间；\n5.根据发明解决问题和专利文本的实施例实验数据和参数，归纳总结【技术效果】。要求字数200-300字之间；\n6.根据专利文本的权利要求1归纳【技术手段】。要求字数200-300字之间；\n7.将技术手段和技术效果以及发明解决问题这3点，提炼出本专利的【发明创新点】。要求字数200-300字之间；\n8.结合发明创新点，以及专利文本的摘要和概述部分，提炼出一段【专利描述】。要求字数200-300字之间；\n9.提炼出本发明的技术实施细节。\n整理格式要求：【专利名称】、【专利申请号】、【申请公布号】、【专利类型】、【法律状态】、【申请人】、【发明人】、【应用领域所属传统行业】、【应用领域所属战略新兴产业】、【解决问题】、【技术效果】、【技术手段】、【创新点】、【专利描述】、【技术实施细节】\n回答只要json格式返回不要缩减内容,key不带【】字符，value的值使用markdown的格式",
                                                  "file_json": [
                                                      {"id": "crr0ako04rj08lqfmcvg", "name": "CN106018500B.PDF",
                                                       "status": "ok", "session_id": 4102,
                                                       "create_time": "2024-09-27T09:07:31.499322+08:00"}]}]})
        print()
        apply_code_list = df['申请号'].str.replace('^CN', '', regex=True).tolist()
        queryset = ZipAnalysisResultModel.objects.filter(patent_info__专利申请号__in=apply_code_list,status='success').all().order_by("id")
        page = self.paginate_queryset(queryset)
        if page is not None:
            serializer = self.get_serializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    def retrieve(self, request, *args, **kwargs):
        serializer = ZipAnalysisResultSerializer(self.get_queryset().filter(zip_id=kwargs['pk']), many=True)
        return Response(serializer.data)
    @action(detail=False, methods=['post'], url_path=ZipAnalysisResultModel._meta.permissions[0][0])
    def upload_zip(self, request):
        user = request.user
        file = request.FILES.get('file')
        file_path = get_file_prefix() + f"/{file.name}"
        unzip_path = file_path.replace('.zip', '')
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        # 打开 ZIP 文件
        unzip(file_path, unzip_path)
        os.remove(file_path)
        all_file_list = []
        for dir, subdir, file_list in os.walk(unzip_path):
            for item_file in file_list:
                if not item_file.startswith('.'):
                    all_file_list.append(item_file)
        total = len(all_file_list)
        user_data_list = ZipAnalysisModel.objects.filter(user_id=user, status='未开始').all()
        all_total = 0
        for i in user_data_list:
            all_total+=i.total

        user_serializer = UserSerializer(user).data
        if (user.all_analysis_count < (total + user_serializer['analysis_count']) or user.all_analysis_count < all_total) and not user.is_superuser:
            raise PermissionDenied(f"可解析数量为{user.all_analysis_count}次,排队中的数据{all_total},已使用{user_serializer['analysis_count']}次")
        zsm = ZipAnalysisModel(name=file.name, status="未开始", total=total,user_id=user)
        zsm.save()
        for dir, subdir, file_list in os.walk(unzip_path):
            for unzip_item_file in file_list:
                if not unzip_item_file.startswith('.'):
                    item_file_path = os.path.join(dir, unzip_item_file)
                    analysis_file.delay(item_file_path, zsm.id, user.id)
        zas = ZipAnalysisSerializer(zsm)
        return Response(zas.data)

    @action(detail=False, methods=['post'], url_path=ZipAnalysisResultModel._meta.permissions[1][0]+"/(?P<zip_analysis_result_id>\d+)",)
    def reanalysis_file(self, request, zip_analysis_result_id):
        zip_result_model = self.get_queryset().filter(id=zip_analysis_result_id).first()
        rest_count = zip_result_model.rest_count
        is_try = request.user.groups.filter(name='试用用户').exists()
        if rest_count >= 5 and is_try:
            return Response({'state': "失败"})
        file_id = zip_result_model.file_id
        zip_id = zip_result_model.zip_id.id
        file_model = FileModel.objects.filter(id=file_id).first()
        zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
        if zip_model.status == "完成":
            zip_model.status = "未开始"
            zip_model.save()
        zip_result_model.delete()
        if is_try:
            rest_count = rest_count+1
        analysis_file.delay(file_model.file_path, zip_id, self.request.user.id, rest_count)
        return Response({'state': "成功"})

    @action(detail=False, methods=['post'],
            url_path="reanalysis_score/(?P<zip_analysis_result_id>\d+)", permission_classes=[IsAdminUser])
    def reanalysis_score(self, request, zip_analysis_result_id):
        zip_result_model = self.get_queryset().filter(id=zip_analysis_result_id).first()
        zip_id = zip_result_model.zip_id.id
        zip_result_model.status='pending'
        zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
        if zip_model.status == "完成":
            zip_model.status = "未开始"
            zip_model.save()
        reanalysis_score.delay(zip_analysis_result_id, zip_id, self.request.user.id,)
        zip_result_model.save()
        return Response({'state': "成功"})
    @action(detail=False, methods=['post'],
            url_path="reanalysis_7/(?P<zip_analysis_result_id>\d+)", permission_classes=[IsAdminUser])
    def reanalysis_7(self, request, zip_analysis_result_id):
        zip_result_model = self.get_queryset().filter(id=zip_analysis_result_id).first()
        zip_id = zip_result_model.zip_id.id
        zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
        zip_result_model.status = 'pending'
        if zip_model.status == "完成":
            zip_model.status = "未开始"
            zip_model.save()
        reanalysis_7.delay(zip_analysis_result_id, zip_id, self.request.user.id,)
        zip_result_model.save()
        return Response({'state': "成功"})
    @action(detail=False, methods=['post'],
            url_path="reanalysis_image/(?P<zip_analysis_result_id>\d+)", permission_classes=[IsAdminUser])
    def reanalysis_image(self, request, zip_analysis_result_id):
        zip_result_model = self.get_queryset().filter(id=zip_analysis_result_id).first()
        zip_id = zip_result_model.zip_id.id
        zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
        zip_result_model.status = 'pending'
        if zip_model.status == "完成":
            zip_model.status = "未开始"
            zip_model.save()
        reanalysis_image.delay(zip_analysis_result_id, zip_id, self.request.user.id,)
        zip_result_model.save()
        return Response({'state': "成功"})
    @action(detail=False, methods=['post'], url_path=ZipAnalysisResultModel._meta.permissions[2][0]+"/(?P<zip_result>\d+)", url_name='down_word')
    def down_word(self, request, zip_result, parent_path=get_file_prefix()):
        zip_result = self.get_queryset().filter(id=zip_result).first()
        if not zip_result:
            raise NotFound(detail="没有数据")
        request_data = request.body
        if request_data:
            request_data = json.loads(request_data.decode("utf-8"))
        file_type = request_data.get("file_type", "docx")
        file_path = down_docx_service(zip_result, file_type, parent_path)
        return FileResponse(open(file_path, 'rb'), as_attachment=True)

    @action(detail=False, methods=['post'], url_path=ZipAnalysisResultModel._meta.permissions[3][0]+"/(?P<zip_id>\d+)")
    @transaction.atomic
    def batch_down_docx(self, request, zip_id):
        request_data = request.body
        if request_data:
            request_data = json.loads(request_data.decode("utf-8"))
        file_type = request_data.get("file_type", "docx")
        zip_analysis = ZipAnalysisModel.objects.get(id=zip_id)
        if file_type == 'docx' and zip_analysis.docx_file:
            return Response(zip_analysis.docx_file.name)
        if file_type == 'pdf' and zip_analysis.pdf_file:
            return Response(zip_analysis.pdf_file.name)
        zip_analysis_result_list = self.get_queryset().filter(zip_id=zip_id).filter(status='success').all()
        if not zip_analysis_result_list:
            raise NotFound(detail="没有可以下载的数据")
        zip_file_name = zip_analysis_result_list.first().zip_id.name
        file_name, ext = os.path.splitext(zip_file_name)
        parent_path = get_file_prefix() + f"{file_name}"
        for zip_analysis_result in zip_analysis_result_list:
            try:
                down_docx_service(zip_analysis_result, file_type, parent_path + "/")
                zip_analysis_result.generate_status = 1
            except Exception as e:
                logger.error("生成docx失败", exc_info=True)
                zip_analysis_result.generate_status = 0
            finally:
                zip_analysis_result.save()
        file_name = get_file_name(zip_id)
        zip_path = os.path.dirname(parent_path) +"/"+file_name+ ".zip"

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for zip_analysis_result in zip_analysis_result_list:
                if file_type == 'docx':
                    zipf.write(zip_analysis_result.docx_file.path, os.path.basename(zip_analysis_result.docx_file.name))
                else:
                    zipf.write(zip_analysis_result.pdf_file.path, os.path.basename(zip_analysis_result.pdf_file.name))
        if os.path.exists(parent_path):
            shutil.rmtree(parent_path)
        zip_analysis.save()
        return Response(zip_path)

    @action(detail=False, methods=['get'],
            url_path="down_zip",permission_classes=[])
    def down_zip(self,request):
        zip_path = self.request.query_params.get("zip_path")
        if zip_path.startswith("upload"):
            return FileResponse(open(zip_path, 'rb'), as_attachment=True)
        return Response(status=404)

    @action(detail=False, methods=['post'], url_path="batch_down_ids")
    def batch_down_ids(self, request):
        request_data = request.body
        if request_data:
            request_data = json.loads(request_data.decode("utf-8"))
        file_type = "docx"
        apply_code_list = df['申请号'].str.replace('^CN', '', regex=True).tolist()
        parent_path = get_file_prefix() + f"zip_id"
        zip_analysis_result_list = []
        for apply_code in apply_code_list:
                zip_analysis_result = ZipAnalysisResultModel.objects.filter(patent_info__专利申请号=apply_code,
                                                      status='success').order_by("-create_time").first()
                if zip_analysis_result:
                    try:
                        zip_analysis_result_list.append(zip_analysis_result)
                        down_docx_service(zip_analysis_result, file_type, parent_path + "/")
                    except Exception as e:
                        logger.error(f"生成docx失败{zip_analysis_result.id}", exc_info=True)
                    finally:
                        zip_analysis_result.save()
        file_name = "zip_file"
        zip_path = os.path.dirname(parent_path) +"/"+file_name+ ".zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for zip_analysis_result in zip_analysis_result_list:
                if file_type == 'docx':
                    zipf.write(zip_analysis_result.docx_file.path, os.path.basename(zip_analysis_result.docx_file.name))
                else:
                    zipf.write(zip_analysis_result.pdf_file.path, os.path.basename(zip_analysis_result.pdf_file.name))
        if os.path.exists(parent_path):
            shutil.rmtree(parent_path)
        return Response(zip_path)

    @action(detail=False, methods=['get'], url_path="get_file_name")
    def get_file_name_request(self,request):
        return Response(get_file_name(request.query_params.get('zip_id'),request.query_params.get('id', None)))


def request_moonshot_ai(data, stream=False):
    message = data['messages']
    problem = get_request_chat(None, message)
    if "response_format" in data:
        response_format = data['response_format']
    else:
        response_format = {"type": "text"}
    model, _ = get_tokens_to_model(problem)
    proxy_request = settings.KIMI_CLIENT.chat.completions.create(
        stream=stream,
        timeout=KIMI_TIMEOUT,
        model=model,
        messages=problem,
        max_tokens=5000,
        response_format=response_format
    )
    return proxy_request


def process_object(obj, level=1):
    markdown = ''
    indent = '  ' * (level - 1)
    for key, value in obj.items():
        key = re.sub("\d+.","-",key)
        if isinstance(value, str):
            markdown += f'{indent}  {key}:{value}\n'
        elif isinstance(value, list):
            markdown += f'{indent}  {key}\n'
            for index, item in enumerate(value):
                if isinstance(item, dict):
                    markdown += process_object(item, level + 1)
                else:
                    markdown += f'{indent}  {item}\n'
        elif isinstance(value, dict):
            try:
                int(key)
                markdown += f'{indent}  {key}{process_object(value, level + 1).strip()}\n'
            except ValueError:
                markdown += f'{indent}  {key}\n{process_object(value, level + 1)}'
    return markdown


def unzip(zip_file, extract_to, encoding='utf-8'):
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        for file_info in zip_ref.infolist():
            try:
                file_info.filename = file_info.filename.encode('cp437').decode('gbk')
            except Exception as e:
                file_info.filename = file_info.filename.encode('utf-8').decode(encoding)
            zip_ref.extract(file_info, extract_to)


def replace_placeholder(doc, data_json, key_ref_link):
    """
    替换文档中的占位符
    """
    use_link_list = []
    for it in doc.paragraphs:
        if it.text.find("{{作画问题}}") != -1:
            it.text = ""
            image_url = data_json.get("作画问题")
            if image_url.strip() != '':
                if "http" in image_url:
                    response = requests.get(data_json.get("作画问题"))
                    if response.status_code == 200:
                        image_name = f'{uuid.uuid4()}.jpg'
                        with open(image_name, 'wb') as f:
                            f.write(response.content)
                        run = it.add_run()
                        run.add_picture(image_name, width=Inches(6.5))
                        os.remove(image_name)
                else:
                        run = it.add_run()
                        run.add_picture(image_url, width=Inches(6.5))
        for key in data_json.keys():
            link_list = key_ref_link.get(key, [])
            if it.text.find("{{" + key + "}}") != -1:
                if '作画问题' == key:
                    continue
                runs = it.runs
                if len(runs) > 1:
                    for run in runs[1:]:
                        r = run._element
                        r.getparent().remove(r)
                        runs[0].text += run.text
                md_data = data_json.get(key)
                if not isinstance(md_data, str):
                    md_data = process_object(md_data)
                md_data = re.sub("([:：。])(\d+)\.\s+", r"\1\n\2.", md_data)
                md_data = re.sub("([:：。])(-)\s+", r"\1\n\2.", md_data)
                md_data = re.sub("【合作企业】|【产业化前景】|行业分类[:：]", "", md_data)
                paragraphs = [i.strip() for i in re.split(r'\n', md_data.rstrip()) if i!='' and i!='\n']
                for index, para in enumerate(paragraphs):
                    break_run = None
                    if key == '合作企业':
                        if re.findall(re_str, para):
                            # text = re.sub(re_str+".*","", text)
                            continue
                    if key in ['产业化前景', '潜在应用场景']:
                        if "未提供" in para or "没有提供" in para or '暂无信息' in para or '：同上' in para:
                            continue
                        # # 正则表达式匹配"XX亿"、"YY亿"、"AA亿"、"BB亿"
                        # pattern = r'\b[A-Z]{1,3}亿\b'
                        # # 使用re.sub进行替换
                        # para = re.sub(pattern, '数亿', para)
                        # pattern = r'\b[A-Z]{1,3}%\b'
                        # # 使用re.sub进行替换
                        # para = re.sub(pattern, '数亿', para)

                    bold_run_list = get_bold_run(key, para)
                    for text, bold in bold_run_list:
                        parts = re.split(r'(\[\d+\])', text)
                        for i in parts:
                            if i.startswith("[") and i.endswith("]"):
                                index = i.replace("[", "").replace("]", "")
                                if index.isdigit():
                                    index = int(index)
                                    if len(link_list) > 0 and len(link_list) > index-1:
                                        url_info = link_list[index - 1]
                                        if url_info['seeMoreUrl'] not in use_link_list:
                                            use_link_list.append(url_info['seeMoreUrl'])
                                        link_index = use_link_list.index(url_info['seeMoreUrl'])
                                        add_hyperlink(it, url_info['seeMoreUrl'], f"[{link_index+1}]")
                                        continue
                                    else:
                                        add_hyperlink(it, "", "")
                                        continue
                            if "{{" + key + "}}" in runs[0].text:
                                runs[0].text = runs[0].text.replace("{{" + key + "}}", i)
                                runs[0].bold = bold
                                break_run = runs[0]
                            else:
                                bold_run = it.add_run(i)
                                bold_run.bold = bold
                                break_run = bold_run
                    if break_run is not None and len(paragraphs) > 1 and index != len(paragraphs) - 1:
                        break_run.add_break()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    find = False
                    for key in data_json.keys():
                        if "{{" + key + "}}" in para.text:
                            find = True
                            para.text = para.text.replace("{{" + key + "}}", str(data_json.get(key)))
                    if not find:
                        if para.text.find("{{") != -1 and para.text.find("}}") != -1:
                            para.text = ""
    use_link_dict = {f"[{index+1}] {value}":value for index,value in enumerate(use_link_list)}
    for it in doc.paragraphs:
        if it.text.find("{{引用链接}}") != -1:
            runs = it.runs
            if len(use_link_dict) == 0:
                for run in runs:
                    r = run._element
                    r.getparent().remove(r)
            else:
                bold_run = it.add_run("数据引用链接")
                bold_run.bold = True
                bold_run.add_break()
                if len(runs) > 0:
                    for run in runs:
                        run.text = ""
                        for key, value in use_link_dict.items():
                            add_hyperlink(it,value,key)
                            it.add_run().add_break()
            if len(use_link_dict.values()) > 0:
                it.add_run().add_break(WD_BREAK.PAGE)

def get_bold_run(key, text):
    new_runs = []
    text = text.replace("（", "(").replace("）",")")
    parts = re.split(r'([\*]{1,2}.*?[\*]{1,2})|(##\s.*?\n)|(###\s.*?\n)|(【\d+\..*?】)|(【.*?】)', text)
    for part in parts:
        if part is not None:
            if part.startswith('**') and part.endswith('**'):
                # 去掉前后的 ** 并加粗
                new_runs.append((part[2:-2], True))
            elif part.startswith('*') and part.endswith('*'):
                # 去掉前后的 ** 并加粗
                new_runs.append((part[1:-1], True))
            elif part.startswith('####'):
                # 去掉前后的 ** 并加粗
                new_runs.append((part[5:].strip(), True))
            elif part.startswith('###'):
                # 去掉前后的 ** 并加粗
                new_runs.append((part[4:].strip(), True))
            elif part.startswith('##'):
                # 去掉前后的 ** 并加粗
                new_runs.append((part[3:].strip(), True))
            elif key in ['最佳应用场景建议', '版本号']:
                new_runs.append((part, True))
            elif re.match(r"^【\d+\.", part):
                new_runs.append((part, True))
            elif re.match(r"^【.*?】", part) and '战略新兴产业分类' not in part:
                new_runs.append((part, True))
            elif re.match(r"^\d+\.$", part):
                new_runs.append((re.sub(r"^\d+\.$","- ", part), False))
            else:
                # 普通文本
                new_runs.append((part, False))
    return new_runs

def add_hyperlink(paragraph, url, text):
    # Create a relationship id
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed attributes
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element
    rPr = OxmlElement('w:rPr')

    # Create a w:rStyle element and set it to 'Hyperlink'
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(ns.qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Create a w:color element and set the color to blue
    color = OxmlElement('w:color')
    color.set(ns.qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    # Create a w:u element and set it to single (underline)
    underline = OxmlElement('w:u')
    underline.set(ns.qn('w:val'), 'single')
    rPr.append(underline)

    # Add the w:rPr element to the w:r element
    new_run.append(rPr)

    # Create a w:t element with the link text
    text_elem = OxmlElement('w:t')
    text_elem.text = text

    # Add the w:t element to the w:r element
    new_run.append(text_elem)

    # Add the w:r element to the w:hyperlink element
    hyperlink.append(new_run)

    # Add the hyperlink element to the paragraph
    paragraph._p.append(hyperlink)


def get_request_chat(session_id, message):
    """
    获取发送给ai的内容
    :return:
    """
    problem = []
    problem += message
    result = []
    for index, item in enumerate(problem):
        if item['role'] == 'user':
            file_json = item.get("file_json", None)
            if file_json is not None:
                for file in file_json:
                    data_list = FileContentModel.objects.filter(file_id=file['id']).all()
                    if data_list:
                        file_content = data_list[0].content
                        if '【具体实施方式】' in file_content:
                            file_content = file_content.split('【具体实施方式】')[0]
                        file_content = file_content
                        result.append({
                            "role": "system",
                            "content": file_content,
                        })
                result.append(item)
            else:
                result.append(item)
        else:
            result.append(item)
    return result


def save_answer(answer, chat_content):
    chat_content_sys = ChatContentModel(role="system", content=answer, session_id=chat_content.session_id,
                                        group_id=chat_content.group_id,
                                        problem_label_id=chat_content.problem_label_id, ref_link=chat_content.ref_link)
    chat_content.save()
    chat_content_sys.save()
@shared_task(bind=True, max_retries=0, default_retry_delay=60)
def analysis_file(self, file_path, zip_id, user_id,rest_count=0):
        """
        解析pdf文件
        :param file_path: 文件地址
        :param zip_id: 压缩包id
        :return:
        """
        user = UserModel.objects.get(id= user_id)
        try:
            zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
            if zip_model is None:
                return
            zip_model.status = "进行中"
            zip_model.save()
            chat_session = ChatSessionModel(user_id=user)
            chat_session.save()
            session_id = chat_session.id
            result_item = ZipAnalysisResultModel(session_id=chat_session, zip_id=ZipAnalysisModel.objects.get(id=zip_id),
                                                 status="success", desc="",user_id=user,rest_count=rest_count)
            result_export = ResultExportModel()
            result_item.name = os.path.basename(file_path)
            try:
                file_model = upload_file_service(session_id, file_path)
                result_item.file_id = file_model['id']
            except Exception as e:
                result_item.desc = "无法提取文本内容"
                result_item.status = "error"
                result_item.save()
                logger.error("无法提取文本内容", exc_info=True)
                return
            problem_list = ProblemLabelModel.objects.order_by("seq").all()
            # 提取的kimi的专利内容整理
            content_extraction = ""
            # 画布信息
            canvas_information = ""
            # 最佳应用场景
            best_application_scenarios = ""
            # 最佳应用市场
            patent_info_str = ""
            tiangong_answer = ""
            answer_json = {}
            for problem in problem_list:
                group_id = str(uuid.uuid4())
                message = problem.name
                session_id = session_id
                chat_content = ChatContentModel(group_id=group_id, role="user",
                                                content=message,
                                                session_id=chat_session, problem_label_id=problem,user_id=user)
                try:
                    seq = problem.seq
                    tiangong_content_list = list(
                        filter(None, [problem.tiangong_content, problem.tiangong_content1, problem.tiangong_content2]))
                    kimi_content_list = [problem.kimi_content, problem.kimi_content1, problem.kimi_content2]
                    image_moonshot_data = {"messages": []}
                    tiangong_data = {"chat_history": [], "stream_resp_type": "all"}
                    # 判断是否需要请求天工
                    for tiangong_content in tiangong_content_list:
                        if tiangong_content is not None and tiangong_content.strip() != '':
                            result_item.status = "success"
                            result_item.desc = ""
                            # 问题7，8，9需要上一个答案作为入参
                            if seq in [7, 8, 9]:
                                request_tiangong_content = content_extraction + "\n" + re.sub("\[\d+\]]","",process_object(answer_json).strip())
                            elif seq in [11]:
                                request_tiangong_content = content_extraction + "\n" + canvas_information
                            elif seq == 14:
                                image_request_content = problem.kimi_content.replace(
                                    "这个场景体现了该专利技术在",
                                    best_application_scenarios + "应用场景")
                                image_moonshot_data['messages'] = [{"role": "user",
                                                                    "content": f"{image_request_content}"}]
                                chat_content.request_kimi = image_moonshot_data
                                # 画图特殊处理
                                request_tiangong_content = ""
                            else:
                                request_tiangong_content = content_extraction
                            tiangong_data['chat_history'] = [
                                {"role": "user", "content": request_tiangong_content + "\n" + tiangong_content}]
                            # 先问天工再将天工的答案给到kimi
                            chat_content.request_tiangong = tiangong_data
                            try:
                                if problem.tiangong_type == 1:
                                    tiangong_answer = request_tiangong_chat_ai(tiangong_data)
                                elif problem.tiangong_type == 2:
                                    tiangong_answer, result_search_list = request_tiangong_copilot_ai(tiangong_data)
                                    chat_content.ref_link = result_search_list
                                elif problem.tiangong_type == 3:
                                    tiangong_answer, result_search_list = request_tiangong_research_ai(tiangong_data)
                                    chat_content.ref_link = result_search_list
                                elif problem.tiangong_type == 4:
                                    # 画图需要先kimi提问处理好再给天工画图
                                    image_content = request_moonshot_ai(image_moonshot_data).choices[0].message.content
                                    tiangong_data['chat_history'] = [
                                        {"role": "user",
                                         "content": f"{tiangong_content}应用场景:{best_application_scenarios}{image_content}"}]
                                    tiangong_answer = request_tiangong_image_ai(tiangong_data)
                                    if tiangong_answer == '':
                                        raise Exception("画图没有结果重新画图:%s", tiangong_data)
                                    image_path = get_file_prefix() + "image.png"
                                    response = requests.get(tiangong_answer)
                                    with open(image_path, 'wb') as f:
                                        f.write(response.content)
                                    tiangong_answer = image_path
                                chat_content.tiangong_answer = tiangong_answer
                                chat_content.request_tiangong = tiangong_data
                                break
                            except Exception as e:
                                logger.error("天工需要重新回答%s", tiangong_answer, exc_info=True)
                                result_item.status = "error"
                                result_item.desc = str(problem.seq) + "问题提问出现错误"
                    else:
                        tiangong_answer = json.dumps(answer_json, ensure_ascii=False)
                    #   如果还是错误直接不往下提问了
                    if result_item.status == 'error':
                        save_answer(answer_json, chat_content)
                        break
                    # 问题14是画图问题，不需要请求kimi所以直接返回
                    if seq == 14:
                        answer_json = {"作画问题": tiangong_answer}
                        save_answer(answer_json, chat_content)
                        continue
                    if seq == 7:
                        answer_json = {"产业化前景": tiangong_answer}
                    special_problem = [6, 7, 9, 11, 15, 16]
                    # 必要内容没有重新提问一次
                    for kimi_content in kimi_content_list:
                        moonshot_data = {"messages": []}
                        if kimi_content is not None and kimi_content.strip() != '':
                            #  每个问题都需要请求kimi
                            # 其他的问题 先问天工再问kimi 问天工带上content_extraction内容
                            # 问kimi的问题5需要跟上专利的基础信息，其余的不需要
                            if seq in [5]:
                                request_kimi = f"{content_extraction}\n本专利的{problem.name}为：{tiangong_answer}\n\n{kimi_content}\n"
                            elif seq == 15:
                                request_kimi = f"{content_extraction}{tiangong_answer}\n{result_item.patent_info.get('专利描述', '')}\n{kimi_content}\n"
                            elif seq in [1]:
                                request_kimi = f"{kimi_content}"
                            elif seq in [12, 13]:
                                request_kimi = f"{content_extraction}{process_object(answer_json)}\n{kimi_content}\n"
                            elif seq == 16:
                                request_kimi = f"{content_extraction}{patent_info_str}\n{kimi_content}\n"
                            else:
                                request_kimi = f"{content_extraction}本专利的{problem.name}为：{tiangong_answer}\n{kimi_content}\n"

                            if seq in special_problem:
                                moonshot_data['messages'] = [{"role": "user",
                                                              "content": request_kimi}]
                            elif seq == 12:
                                moonshot_data['messages'] = [{"role": "user",
                                                              "content": f"{request_kimi}\n"+
                                                                         """请使用如下JSON格式输出你的回复:{"权利要求的清晰性":"权利要求的清晰性分数","保护范围的深度":"保护范围的深度分数","受到技术挑战的可能性":"受到技术挑战的可能性分数","技术创新性":"技术创新性分数","技术独特性":"技术独特性分数","技术复杂性":"技术复杂性分数","资源需求":"资源需求分数","技术转化周期":"技术转化周期分数","风险与不确定性":"风险与不确定性分数","社会与经济影响":"社会与经济影响分数","潜在市场规模":"潜在市场规模分数","竞争格局":"竞争格局分数","预期利润":"预期利润分数","市场份额":"市场份额分数","渠道拓展与维护难度":"渠道拓展与维护难度分数","营销资源投入":"营销资源投入分数"}"""}]
                                moonshot_data['response_format'] = {"type": "json_object"}
                            else:
                                moonshot_data['messages'] = [{"role": "user",
                                                              "content": f"{request_kimi}\n"+
                                                                         f"回答只要json格式返回不要缩减内容,key不带【】字符，value的值使用markdown的格式"}]
                            if seq in [1, 12, 13]:
                                moonshot_data['messages'][0]['file_json'] = []
                                moonshot_data['messages'][0]['file_json'].append(file_model)
                            if seq == 16:
                                file_model_2018 = upload_file_service(session_id,
                                                                      '战略性新兴产业分类(2018).doc')
                                moonshot_data['messages'][0]['file_json'] = [file_model_2018]
                            chat_content.request_kimi = moonshot_data
                            proxy_request = request_moonshot_ai(moonshot_data)
                            result_item.desc = ""
                            result_item.status = 'success'
                            try:
                                answer_json = {}
                                moonshot_answer = proxy_request.choices[0].message.content
                                if seq in special_problem:
                                    answer_json = {problem.name: moonshot_answer}
                                else:
                                    answer_json = json.loads(moonshot_answer.replace(" “",'"').replace("”",'”').replace("`", "").replace("json", ""),
                                                             strict=False)
                                # 获取该任务需要的key
                                problem_key_list = PROBLEM_MAPPING_KEY_JSON.get(str(seq), set())
                                keys = answer_json.keys()
                                # 判断是否必要的key是否都存在
                                if len(problem_key_list) > 0 and len(problem_key_list - keys) > 0:
                                    raise Exception(f"缺少必要的数据{problem_key_list - keys}")
                                score_list = []
                                if seq == 12:
                                    for key in answer_json.keys():
                                        dict_item = answer_json.get(key, "")
                                        if isinstance(dict_item, int):
                                            score_list.append(answer_json.get(key))
                                        if isinstance(dict_item, str) and str.isdigit(dict_item):
                                            score_list.append(int(answer_json.get(key)))
                                        if isinstance(dict_item, dict):
                                            for value in dict_item.values():
                                                if isinstance(value, int):
                                                    score_list.append(value)
                                    if len(score_list) != 16:
                                        raise Exception(f"评分数量不为16")
                                break
                            except Exception as e:
                                # 需要重新回答
                                logger.error("kimi需要重新回答:%s", str(proxy_request), exc_info=True)
                                result_item.status = "error"
                                result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误"

                    #   如果还是错误直接不往下提问了
                    if result_item.status == 'error':
                        break
                    if seq == 10:
                        canvas_information = process_object(answer_json)
                    if seq == 8:
                        # 提取最佳应用场景
                        best_application_scenarios = answer_json.get(
                            "最佳应用场景建议", answer_json.get("专利名称", answer_json.get("专利描述", "")))
                    if seq == 1:
                        content_extraction += (
                            f"本专利信息如下：\n应用领域所属传统行业:{answer_json.get('应用领域所属传统行业', '')}"
                            f"\n应用领域所属战略新兴产业:{answer_json.get('应用领域所属战略新兴产业', '')}"
                            f"\n解决问题:{answer_json.get('解决问题', '')}"
                            f"\n技术效果:{answer_json.get('技术效果', '')}"
                            f"\n技术手段:{answer_json.get('技术手段', '')}"
                            f"\n创新点:{answer_json.get('创新点', '')}"
                            f"\n专利描述:{answer_json.get('专利描述', '')}")
                        apply_code = answer_json.get("专利申请号").strip().replace(" ", "")
                        # 存在的话更新数据
                        old_result_export = ResultExportModel.objects.filter(apply_code=apply_code).first()
                        if old_result_export:
                            result_export = old_result_export
                        result_export.apply_code = answer_json.get("专利申请号").strip().replace(" ","")
                        result_export.patent_name = answer_json.get("专利名称")
                        patent_info_str = "【专利描述】" + answer_json.get('专利描述', '')
                        inventor_str = answer_json.get("发明人")
                        if isinstance(inventor_str, list):
                            inventor_list = inventor_str
                            answer_json['发明人'] = ','.join(inventor_list)
                        else:
                            inventor_list = [i.strip() for i in re.split('[;、,；，\s]', inventor_str)]
                        result_export.department = ",".join(
                            [u.department for u in UserModel.objects.filter(username__in=inventor_list)])
                        result_item.patent_info = answer_json
                    if seq == 7:
                        result_export.ctp = answer_json.get("产业化前景","").replace("#","")
                    if seq == 5:
                        doip = answer_json.get("本发明与现有技术的比较分析")
                        if isinstance(doip, dict):
                            doip = process_object(doip)
                        result_export.doip = doip
                    if seq == 11:
                        data_markdown = answer_json.get("合作企业")
                        data = [i.strip() for i in data_markdown.split("\n") if i != '']
                        c = []
                        d = set()
                        for line in data:
                            if re.match(r'^\d.', line):
                                line = line.replace("*", "")
                                d_item = re.findall(r'[（\(](.*?市)[\)）]', line)
                                if d_item:
                                    d.add(d_item[0])
                                d_item = re.findall(r'-\s*(.*?市)', line)
                                if d_item:
                                    d.add(d_item[0])
                                c.append(line)
                            if re.match(re_str, line):
                                d.add(re.sub(re_str, "", line))
                        data = [i for i in zip_longest(c, d, fillvalue='')]
                        result_export.cooperative_enterprises = ", ".join([i[0] for i in data])
                        result_export.cooperative_city = ",".join([i[1] for i in data if i[1] != ''])
                    save_answer(answer_json, chat_content)
                except Exception as e:
                    result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误\n"
                    result_item.status = 'error'
                    logger.error("错误:答案天工:%s,kimi%s", tiangong_answer, json.dumps(answer_json, ensure_ascii=False),
                                 exc_info=True)
                    save_answer({}, chat_content)
                    break
            result_item.save()
            result_export.zip_analysis_id = result_item
            if result_item.status == 'success':
                result_export.user_id = user
                result_export.save()
            # 加锁
            lock_zip_model_list = ZipAnalysisModel.objects.select_for_update().filter(id=zip_id)
            with transaction.atomic():
                for lock_zip_model in lock_zip_model_list:
                    zip_analysis_result_count = ZipAnalysisResultModel.objects.select_for_update().filter(zip_id=zip_id).count()
                    if int(zip_analysis_result_count) == int(lock_zip_model.total):
                        lock_zip_model.status = "完成"
                        lock_zip_model.save()
        except Exception as e:
            logger.error("出现错误", exc_info=True)
            raise e
        finally:
            # 加锁
            lock_zip_model_list = ZipAnalysisModel.objects.select_for_update().filter(id=zip_id)
            with transaction.atomic():
                for lock_zip_model in lock_zip_model_list:
                    zip_analysis_result_count = ZipAnalysisResultModel.objects.select_for_update().filter(zip_id=zip_id).count()
                    if int(zip_analysis_result_count) == int(lock_zip_model.total):
                        lock_zip_model.status = "完成"
                        lock_zip_model.save()

@shared_task(bind=True, max_retries=0, default_retry_delay=60)
def reanalysis_score(self, zip_result_id, zip_id, user_id):
        """
        解析pdf文件
        :param zip_id: 压缩包id
        :return:
        """
        problem_id = 12
        user = UserModel.objects.get(id= user_id)
        try:
            zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
            if zip_model is None:
                return
            zip_model.status = "进行中"
            zip_model.save()
            result_item = ZipAnalysisResultModel.objects.get(id=zip_result_id)
            result_item.status = 'doing'
            result_item.save()
            problem_list = ProblemLabelModel.objects.filter(id=problem_id).order_by("seq").all()
            ChatContentModel.objects.filter(problem_label_id=problem_id, session_id=result_item.session_id.id).delete()
            # 提取的kimi的专利内容整理
            tiangong_answer = ""
            answer_json = {}
            file_model = FileSerializer(FileModel.objects.get(id=result_item.file_id))
            for problem in problem_list:
                group_id = str(uuid.uuid4())
                message = problem.name
                chat_content = ChatContentModel(group_id=group_id, role="user",
                                                content=message,
                                                session_id=result_item.session_id, problem_label_id=problem,user_id=user)
                try:
                    seq = problem.seq
                    kimi_content_list = [problem.kimi_content, problem.kimi_content1, problem.kimi_content2]
                    for kimi_content in kimi_content_list:
                        moonshot_data = {"messages": []}
                        if kimi_content is not None and kimi_content.strip() != '':
                            moonshot_data['messages'] = [{"role": "user",
                                                          "content": f"{kimi_content}\n"+
                                                                     """请使用如下JSON格式输出你的回复:{"权利要求的清晰性":"权利要求的清晰性分数","保护范围的深度":"保护范围的深度分数","受到技术挑战的可能性":"受到技术挑战的可能性分数","技术创新性":"技术创新性分数","技术独特性":"技术独特性分数","技术复杂性":"技术复杂性分数","资源需求":"资源需求分数","技术转化周期":"技术转化周期分数","风险与不确定性":"风险与不确定性分数","社会与经济影响":"社会与经济影响分数","潜在市场规模":"潜在市场规模分数","竞争格局":"竞争格局分数","预期利润":"预期利润分数","市场份额":"市场份额分数","渠道拓展与维护难度":"渠道拓展与维护难度分数","营销资源投入":"营销资源投入分数"}"""}]
                            moonshot_data['response_format'] = {"type": "json_object"}
                            moonshot_data['messages'][0]['file_json'] = [file_model.data]
                            chat_content.request_kimi = moonshot_data
                            proxy_request = request_moonshot_ai(moonshot_data)
                            result_item.desc = ""
                            result_item.status = 'success'
                            try:
                                answer_json = {}
                                moonshot_answer = proxy_request.choices[0].message.content
                                answer_json = json.loads(moonshot_answer.replace(" “",'"').replace("”",'”').replace("`", "").replace("json", ""),
                                                             strict=False)
                                # 获取该任务需要的key
                                problem_key_list = PROBLEM_MAPPING_KEY_JSON.get(str(seq), set())
                                keys = answer_json.keys()
                                # 判断是否必要的key是否都存在
                                if len(problem_key_list) > 0 and len(problem_key_list - keys) > 0:
                                    raise Exception(f"缺少必要的数据{problem_key_list - keys}")
                                score_list = []
                                if seq == 12:
                                    for key in answer_json.keys():
                                        dict_item = answer_json.get(key, "")
                                        if isinstance(dict_item, int):
                                            score_list.append(answer_json.get(key))
                                        if isinstance(dict_item, dict):
                                            for value in dict_item.values():
                                                if isinstance(value, int):
                                                    score_list.append(value)
                                    if len(score_list) != 16:
                                        raise Exception(f"评分数量不为16")
                                break
                            except Exception as e:
                                # 需要重新回答
                                logger.error("kimi需要重新回答:%s", str(proxy_request), exc_info=True)
                                result_item.status = "error"
                                result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误"

                    #   如果还是错误直接不往下提问了
                    if result_item.status == 'error':
                        break
                    save_answer(answer_json, chat_content)
                except Exception as e:
                    result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误\n"
                    result_item.status = 'error'
                    logger.error("错误:答案天工:%s,kimi%s", tiangong_answer, json.dumps(answer_json, ensure_ascii=False),
                                 exc_info=True)
                    save_answer({}, chat_content)
                    break
            result_item.status = 'success'
            result_item.save()
        except Exception as e:
            logger.error("出现错误", exc_info=True)
            raise e
        finally:            # 加锁
            lock_zip_model_list = ZipAnalysisModel.objects.select_for_update().filter(id=zip_id)
            with transaction.atomic():
                for lock_zip_model in lock_zip_model_list:
                    zip_analysis_result_count = ZipAnalysisResultModel.objects.select_for_update().filter(zip_id=zip_id).count()
                    if int(zip_analysis_result_count) == int(lock_zip_model.total):
                        lock_zip_model.status = "完成"
                        lock_zip_model.save()


@shared_task(bind=True, max_retries=0, default_retry_delay=60)
def reanalysis_7(self, zip_result_id, zip_id, user_id):
        """
        解析pdf文件
        :param zip_id: 压缩包id
        :return:
        """
        problem_id = [6,7,8]
        user = UserModel.objects.get(id= user_id)
        try:
            zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
            if zip_model is None:
                return
            zip_model.status = "进行中"
            zip_model.save()
            result_item = ZipAnalysisResultModel.objects.get(id=zip_result_id)
            result_item.status = 'doing'
            result_item.save()
            problem_list = ProblemLabelModel.objects.filter(id__in=problem_id).order_by("seq").all()
            ChatContentModel.objects.filter(problem_label_id__in = problem_id, session_id=result_item.session_id.id).delete()
            # 提取的kimi的专利内容整理
            tiangong_answer = ""
            answer_json = {}
            file_model = FileSerializer(FileModel.objects.get(id=result_item.file_id))
            for problem in problem_list:
                seq = problem.seq
                group_id = str(uuid.uuid4())
                message = problem.name
                chat_content = ChatContentModel(group_id=group_id, role="user",
                                                content=message,
                                                session_id=result_item.session_id, problem_label_id=problem,user_id=user)
                try:
                    kimi_content_list = list(filter(lambda x: x not in (None, ""),
                                    [problem.kimi_content, problem.kimi_content1, problem.kimi_content2]))
                    for kimi_content in kimi_content_list:
                        moonshot_data = {"messages": [{"role":i.role,"content":json.dumps(i.content,ensure_ascii=False)} for i in ChatContentModel.objects.filter(problem_label_id__in=problem_id,
                                                        session_id=result_item.session_id.id).all()]}
                        if kimi_content is not None and kimi_content.strip() != '':
                            if seq == 8:
                                kimi_content+="""回答按一下格式返回:{"分析逻辑和理由": "分析逻辑和理由", "最佳应用场景建议": "最佳应用场景建议"}，value的值使用markdown的格式"""
                                moonshot_data['response_format'] = {"type": "json_object"}
                            moonshot_data['messages'].append({"role": "user",
                                                          "content": f"{kimi_content}\n"})
                            moonshot_data['messages'][0]['file_json'] = [file_model.data]
                            chat_content.request_kimi = moonshot_data
                            proxy_request = request_moonshot_ai(moonshot_data)
                            result_item.desc = ""
                            result_item.status = 'success'
                            try:
                                moonshot_content = proxy_request.choices[0].message.content
                                if seq == 7:
                                    answer_json = {"产业化前景":moonshot_content}
                                elif seq == 6:
                                    answer_json = {"潜在应用场景": moonshot_content}
                                else:
                                    answer_json = json.loads(
                                        moonshot_content.replace(" “", '"').replace("”", '”').replace("`", "").replace(
                                            "json", ""),
                                        strict=False)
                                break
                            except Exception as e:
                                # 需要重新回答
                                logger.error("kimi需要重新回答:%s", str(proxy_request), exc_info=True)
                                result_item.status = "error"
                                result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误"

                    #   如果还是错误直接不往下提问了
                    if result_item.status == 'error':
                        break
                    save_answer(answer_json, chat_content)
                except Exception as e:
                    result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误\n"
                    result_item.status = 'error'
                    logger.error("错误:答案天工:%s,kimi%s", tiangong_answer, json.dumps(answer_json, ensure_ascii=False),
                                 exc_info=True)
                    save_answer({}, chat_content)
                    break
            result_item.status = 'success'
            result_item.save()
        except Exception as e:
            logger.error("出现错误", exc_info=True)
            raise e
        finally:            # 加锁
            lock_zip_model_list = ZipAnalysisModel.objects.select_for_update().filter(id=zip_id)
            with transaction.atomic():
                for lock_zip_model in lock_zip_model_list:
                    zip_analysis_result_count = ZipAnalysisResultModel.objects.select_for_update().filter(zip_id=zip_id).count()
                    if int(zip_analysis_result_count) == int(lock_zip_model.total):
                        lock_zip_model.status = "完成"
                        lock_zip_model.save()

@shared_task(bind=True, max_retries=0, default_retry_delay=60)
def reanalysis_image(self, zip_result_id, zip_id, user_id):
    """
    解析pdf文件
    :param zip_id: 压缩包id
    :return:
    """
    problem_id = [14]
    user = UserModel.objects.get(id=user_id)
    try:
        zip_model = ZipAnalysisModel.objects.filter(id=zip_id).first()
        if zip_model is None:
            return
        zip_model.status = "进行中"
        zip_model.save()
        result_item = ZipAnalysisResultModel.objects.get(id=zip_result_id)
        result_item.status = 'doing'
        result_item.save()
        problem_list = ProblemLabelModel.objects.filter(id__in=problem_id).order_by("seq").all()
        ChatContentModel.objects.filter(problem_label_id__in=problem_id, session_id=result_item.session_id.id).delete()
        # 提取的kimi的专利内容整理
        tiangong_answer = ""
        answer_json = {}
        for problem in problem_list:
            best_application_scenarios = ChatContentModel.objects.filter(problem_label_id=6, role='system',session_id=result_item.session_id).first().content
            group_id = str(uuid.uuid4())
            message = problem.name
            chat_content = ChatContentModel(group_id=group_id, role="user",
                                            content=message,
                                            session_id=result_item.session_id, problem_label_id=problem, user_id=user)
            try:
                kimi_content_list = list(filter(lambda x: x not in (None, ""),
                                                [problem.kimi_content, problem.kimi_content1, problem.kimi_content2]))
                for kimi_content in kimi_content_list:
                    image_moonshot_data = {"messages": []}
                    tiangong_data = {"chat_history": [], "stream_resp_type": "all"}
                    if kimi_content is not None and kimi_content.strip() != '':
                        try:
                            image_request_content = problem.kimi_content.replace(
                                "这个场景体现了该专利技术在",
                                best_application_scenarios.get('潜在应用场景') + "应用场景")
                            image_moonshot_data['messages'] = [{"role": "user",
                                                                "content": f"{image_request_content}"}]
                            # 画图需要先kimi提问处理好再给天工画图
                            image_content = request_moonshot_ai(image_moonshot_data).choices[0].message.content
                            tiangong_data['chat_history'] = [
                                {"role": "user",
                                 "content": f"{image_content}"}]
                            tiangong_answer = request_tiangong_image_ai(tiangong_data)
                            if tiangong_answer == '':
                                raise Exception("画图没有结果重新画图:%s", tiangong_data)
                            image_path = get_file_prefix()+"image.png"
                            response = requests.get(tiangong_answer)
                            with open(image_path, 'wb') as f:
                                f.write(response.content)
                            chat_content.content = "作画问题"
                            answer_json = {"作画问题": image_path}
                            chat_content.request_tiangong = tiangong_data
                            break
                        except Exception as e:
                            # 需要重新回答
                            logger.error("kimi需要重新回答:%s", exc_info=True)
                            result_item.status = "error"
                            result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误"

                #   如果还是错误直接不往下提问了
                if result_item.status == 'error':
                    break
                save_answer(answer_json, chat_content)
            except  Exception as e:
                result_item.desc = result_item.desc + "\n" + str(problem.seq) + "问题提问出现错误\n"
                result_item.status = 'error'
                logger.error("错误:答案天工:%s,kimi%s", tiangong_answer, json.dumps(answer_json, ensure_ascii=False),
                             exc_info=True)
                save_answer({}, chat_content)
                break
        result_item.status = 'success'
        result_item.save()
    except Exception as e:
        logger.error("出现错误", exc_info=True)
        raise e
    finally:  # 加锁
        lock_zip_model_list = ZipAnalysisModel.objects.select_for_update().filter(id=zip_id)
        with transaction.atomic():
            for lock_zip_model in lock_zip_model_list:
                zip_analysis_result_count = ZipAnalysisResultModel.objects.select_for_update().filter(
                    zip_id=zip_id).count()
                if int(zip_analysis_result_count) == int(lock_zip_model.total):
                    lock_zip_model.status = "完成"
                    lock_zip_model.save()


@transaction.atomic
def down_docx_service(zip_analysis_result, file_type, parent_path=get_file_prefix()):
    chat_list = ChatContentModel.objects.filter(session_id=zip_analysis_result.session_id.id).filter(role='system').order_by('id').all()
    if file_type == 'docx' and zip_analysis_result.docx_file:
        return zip_analysis_result.docx_file.path
    elif file_type == 'pdf' and zip_analysis_result.pdf_file:
        return zip_analysis_result.pdf_file.path
    else:
        formatted_number = f"{VERSION_PREFIX}{zip_analysis_result.id:04}"
        file_name = get_file_name(zip_id=None,zip_result_id=zip_analysis_result.id)
        content_json = {"版本号": formatted_number}
        key_ref_link = {}
        for chat in chat_list:
            ref_link = chat.ref_link
            content_item = chat.content
            if ref_link:
                key_ref_link = key_ref_link | {key: ref_link for key in content_item.keys()}
            content_json = content_json | content_item
            if chat.problem_label_id.seq == 14:
                if 'http' in chat.content.get("作画问题"):
                    image_path = get_file_prefix()+"image.png"
                    response = requests.get(chat.content.get("作画问题"))
                    with open(image_path, 'wb') as f:
                        f.write(response.content)
                    chat.content = {"作画问题":image_path}
                    chat.save()
            if chat.problem_label_id.seq == 12:
                content_json.update(get_score(zip_analysis_result.id, content_json))
                # content_item['申请公布号'] = content_json['申请公布号']
                # content_item = {k: int(v) if isinstance(v,str) and v.isdigit() else v for k, v in content_item.items()}
                # compute_score(content_item, zip_analysis_result)
                # content_json['专利保护范围分数'] = (Decimal(content_item['权利要求的清晰性'] * 0.1 + content_item.get("保护范围的深度",content_item.get("专利对核心技术细节的保护程度")) * 0.1
                #                                             +content_item['说明书分数'] * 0.35 + content_item['独立权利要求分数'] * 0.45
                #                                             )
                #                                     .quantize(Decimal('1.0'),
                #                                                                                                rounding=ROUND_HALF_UP))
                # content_json['专利稳定性分数'] = (Decimal(content_item['受到技术挑战的可能性']*0.1 + content_item['专利类型得分']*0.25+content_item['简单同族分数']*0.4+content_item['法律状态得分']*0.25)
                #                                   .quantize(Decimal('1.0'), rounding=ROUND_HALF_UP))
                # content_json['技术先进性分数'] = Decimal(content_item['技术创新性'] * 0.1 + content_item['技术独特性'] * 0.1 + content_item['被引用分数'] * 0.8).quantize(Decimal('1.0'),
                #                                                                                              rounding=ROUND_HALF_UP)
                # content_json['技术实施难度分数'] = Decimal(content_item['技术复杂性'] * 0.3 + content_item['资源需求'] * 0.3 + content_item['技术转化周期'] * 0.2 + \
                #                                            content_item[
                #                                                '风险与不确定性'] * 0.2).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
                # content_json['社会与经济影响分数'] = Decimal(content_item['社会与经济影响']*0.2+content_item['战略新兴产业分类得分']*0.8).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
                # content_json['市场需求分数'] = Decimal(content_item['潜在市场规模'] * 0.2 + content_item['竞争格局'] * 0.8).quantize(Decimal('1.0'),
                #                                                                                              rounding=ROUND_HALF_UP)
                # content_json['潜在商业价值分数'] = Decimal(content_item['预期利润'] * 0.05 + content_item['市场份额'] * 0.05 + content_item['年限分数']*0.65 + content_item['当前申请人分数']*0.25).quantize(Decimal('1.0'),
                #                                                                                                  rounding=ROUND_HALF_UP)
                # content_json['市场推广难度分数'] = Decimal(content_item['渠道拓展与维护难度'] * 0.7 + content_item['营销资源投入'] * 0.3).quantize(Decimal('1.0'),
                #                                                                                                  rounding=ROUND_HALF_UP)
                # content_json['综合评分'] = Decimal(
                #     (content_json['专利保护范围分数'] * Decimal(0.6) + content_json['专利稳定性分数'] * Decimal(0.4)) * Decimal(
                #         0.4) + (
                #             content_json['技术先进性分数'] * Decimal(0.6) + content_json['技术实施难度分数'] * Decimal(0.1) +
                #             content_json[
                #                 '社会与经济影响分数'] * Decimal(0.3)) * Decimal(0.4) + (
                #             content_json['市场需求分数'] * Decimal(0.15) + content_json['潜在商业价值分数'] * Decimal(0.7) +
                #             content_json[
                #                 '市场推广难度分数'] * Decimal(0.15)
                #     ) * Decimal(0.2)).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
        temp_data_json = {key: "" for key in TEMP_KEY_NAME_LIST}
        for key in content_json.keys():
            if key in TEMP_KEY_NAME_LIST:
                if isinstance(content_json.get(key), dict):
                    temp_data_json[key] = process_object(content_json.get(key))
                elif isinstance(content_json.get(key), list):
                    temp_data_json[key] = "\n".join(content_json.get(key))
                else:
                    temp_data_json[key] = content_json.get(key)
        p_code = temp_data_json.get("申请公布号","").replace(" ","")[:-1]
        # df = pd.read_excel('index.xlsx')
        # df = df.fillna("无")
        patent_info = df[df['公开(公告)号'] == p_code+"B"]
        if patent_info.empty:
            patent_info = df[df['公开(公告)号'] == p_code + "A"]
        if patent_info.empty:
            patent_info = df[df['公开(公告)号'] == p_code + "C"]
        if patent_info.empty:
            patent_info = df[df['公开(公告)号'] == p_code + "U"]
        if patent_info.empty:
            patent_info = df[df['公开(公告)号'] == p_code + "S"]
        filtered_data1 = patent_info['一级产业分类']
        filtered_data2 = patent_info['二级产业分类']
        temp_data_json['行业分类'] = "一级产业分类："+str(filtered_data1.iloc[0] if not filtered_data1.empty
            else '无')+"\n"+"二级产业分类："+str(filtered_data2.iloc[0] if not filtered_data2.empty else '无')
        clazz_type_data = get_clazz_type(p_code)
        if not clazz_type_data:
            temp_data_json['国民经济分类'] = "国民经济行业分类\n无"
        else:
            temp_data_json['国民经济分类'] = ""
        for clazz_index, value in enumerate(clazz_type_data.values()):
            if len(clazz_type_data)>1:
                temp_data_json['国民经济分类']+=f"国民经济行业分类{clazz_index+1}\n"
            else:
                temp_data_json['国民经济分类'] += "国民经济行业分类\n"
            category_str = "；".join(value['category_list'])
            temp_data_json['国民经济分类']+=f"门类：{value['big_name']}\n大类：{category_str}\n"
        doc = Document('概念验证报告.docx')
        replace_placeholder(doc, temp_data_json, key_ref_link)
        file_path_docx = parent_path + file_name + ".docx"
        down_file_path = parent_path + file_name + ".pdf"
        os.makedirs(os.path.dirname(down_file_path), exist_ok=True)
        doc.save(file_path_docx)
        with open(file_path_docx, 'rb') as f:
            zip_analysis_result.docx_file.save(file_name+'.docx', File(f), save=True)
        if file_type == 'docx':
            down_file_path = file_path_docx
        else:
            # 使用 unoconv 将 docx 转换为 pdf
            subprocess.run(
                [settings.LIBREOFFICE, "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(down_file_path),
                 file_path_docx])
            os.remove(file_path_docx)
            with open(down_file_path, 'rb') as f:
                zip_analysis_result.pdf_file.save(file_name+'.pdf', File(f), save=True)
        zip_analysis_result.save()
        return down_file_path
def get_file_name(zip_id, zip_result_id=None):
        if zip_result_id:
            data = ZipAnalysisResultModel.objects.filter(id=zip_result_id)
        else:
            data = ZipAnalysisResultModel.objects.filter(zip_id=zip_id)
        data = data.filter(status='success')
        if len(data) > 0:
            first = data[0]
            patent_name = data[0].patent_info['专利名称']
            count = len(data)
            formatted_number = f"{VERSION_PREFIX}{first.id:04}"
            if len(data) == 1:
                file_name = f'{formatted_number}-{patent_name}'
            else:
                file_name = f'{formatted_number}等{count}件专利-{patent_name}等{count}件专利'
        else:
            raise NotFound(detail="没有数据")
        file_name = file_name.replace("/","")
        return file_name

def get_score(id, content_item):
    zip_analysis_result = ZipAnalysisResultModel.objects.get(pk=id)
    content_json = {}
    content_item['申请公布号'] = zip_analysis_result.patent_info['申请公布号']
    content_item = {k: int(v) if isinstance(v, str) and v.isdigit() else v for k, v in content_item.items()}
    compute_score(content_item, zip_analysis_result)
    content_json['专利保护范围分数'] = (Decimal(
        content_item['权利要求的清晰性'] * 0.1 + content_item.get("保护范围的深度", content_item.get(
            "专利对核心技术细节的保护程度")) * 0.1
        + content_item['说明书分数'] * 0.35 + content_item['独立权利要求分数'] * 0.45
        )
                                        .quantize(Decimal('1.0'),
                                                  rounding=ROUND_HALF_UP))
    content_json['专利稳定性分数'] = (Decimal(
        content_item['受到技术挑战的可能性'] * 0.1 + content_item['专利类型得分'] * 0.25 + content_item[
            '简单同族分数'] * 0.4 + content_item['法律状态得分'] * 0.25)
                                      .quantize(Decimal('1.0'), rounding=ROUND_HALF_UP))
    content_json['技术先进性分数'] = Decimal(
        content_item['技术创新性'] * 0.1 + content_item['技术独特性'] * 0.1 + content_item[
            '被引用分数'] * 0.8).quantize(Decimal('1.0'),
                                          rounding=ROUND_HALF_UP)
    content_json['技术实施难度分数'] = Decimal(
        content_item['技术复杂性'] * 0.3 + content_item['资源需求'] * 0.3 + content_item['技术转化周期'] * 0.2 + \
        content_item[
            '风险与不确定性'] * 0.2).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
    content_json['社会与经济影响分数'] = Decimal(
        content_item['社会与经济影响'] * 0.2 + content_item['战略新兴产业分类得分'] * 0.8).quantize(Decimal('1.0'),
                                                                                                    rounding=ROUND_HALF_UP)
    content_json['市场需求分数'] = Decimal(
        content_item['潜在市场规模'] * 0.2 + content_item['竞争格局'] * 0.8).quantize(Decimal('1.0'),
                                                                                      rounding=ROUND_HALF_UP)
    content_json['潜在商业价值分数'] = Decimal(
        content_item['预期利润'] * 0.05 + content_item['市场份额'] * 0.05 + content_item['年限分数'] * 0.65 +
        content_item['当前申请人分数'] * 0.25).quantize(Decimal('1.0'),
                                                        rounding=ROUND_HALF_UP)
    content_json['市场推广难度分数'] = Decimal(
        content_item['渠道拓展与维护难度'] * 0.7 + content_item['营销资源投入'] * 0.3).quantize(Decimal('1.0'),
                                                                                                rounding=ROUND_HALF_UP)
    law_score = (content_json['专利保护范围分数'] * Decimal(0.6) + content_json['专利稳定性分数'] * Decimal(0.4)).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
    tech_score = (content_json['技术先进性分数'] * Decimal(0.6) + content_json['技术实施难度分数'] * Decimal(0.1) +
            content_json[
                '社会与经济影响分数'] * Decimal(0.3)).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
    market_score = (
            content_json['市场需求分数'] * Decimal(0.15) + content_json['潜在商业价值分数'] * Decimal(0.7) +
            content_json[
                '市场推广难度分数'] * Decimal(0.15)
    ).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
    content_json['法律得分'] = law_score
    content_json['市场得分'] = market_score
    content_json['技术得分'] = tech_score
    content_json['综合评分'] = Decimal(
        law_score* Decimal(
        0.4) + tech_score * Decimal(0.4) + market_score * Decimal(0.2)).quantize(Decimal('1.0'), rounding=ROUND_HALF_UP)
    return content_json

def compute_score(temp_data_json,zarm):
    """增加几个评分信息"""
    p_code = temp_data_json.get("申请公布号", "").replace(" ", "")[:-1]
    patent_info = df[df['公开(公告)号'] == p_code + "B"]
    if patent_info.empty:
        patent_info = df[df['公开(公告)号'] == p_code + "A"]
    if patent_info.empty:
        patent_info = df[df['公开(公告)号'] == p_code + "C"]
    if patent_info.empty:
        patent_info = df[df['公开(公告)号'] == p_code + "U"]
    if patent_info.empty:
        patent_info = df[df['公开(公告)号'] == p_code + "S"]
    ref_score = 60
    legal_score = 60
    patent_type_score = 30
    year_score = 40
    temp_data_json['战略新兴产业分类得分'] = 45
    temp_data_json['被引用分数'] = ref_score
    temp_data_json['当前申请人分数'] = 60
    temp_data_json['简单同族分数'] = 60
    temp_data_json['独立权利要求分数'] = 55
    temp_data_json['法律状态得分'] = legal_score
    temp_data_json['专利类型得分'] = patent_type_score
    temp_data_json['年限分数'] = year_score
    temp_data_json['说明书分数'] = 50
    if not patent_info.empty:
        # classify_data1 = patent_info['一级产业分类']
        # classify_data2 = patent_info['二级产业分类']
        # temp_data_json['行业分类'] = "一级产业分类：" + str(classify_data1.iloc[0] if not classify_data1.empty
        #                                                    else '无') + "\n" + "二级产业分类：" + str(
        #     classify_data2.iloc[0] if not classify_data2.empty else '无')
        temp_data_json['战略新兴产业分类得分'] = 45
        if not patent_info['战略新兴产业分类'].empty and patent_info['战略新兴产业分类'].iloc[0]!='-':
            temp_data_json['战略新兴产业分类得分'] = 90
        ref_count = patent_info['被引用专利数量'].iloc[0]
        if ref_count>15:
            ref_score = 100
        elif ref_count>10 and ref_count< 14:
            ref_score = 90
        elif ref_count>5 and ref_count< 9:
            ref_score = 80
        elif ref_count>1 and ref_count< 4:
            ref_score = 70
        temp_data_json['被引用分数'] = ref_score
        application_count = patent_info['[标]当前申请(专利权)人'].iloc[0].count('|') + 1
        if application_count == 1:
            temp_data_json['当前申请人分数'] = 90
        if application_count == 2:
            temp_data_json['当前申请人分数'] = 80
        temp_data_json['简单同族分数'] = min(60+(patent_info['简单同族'].iloc[0].count('|')+1)*10,100)
        temp_data_json['独立权利要求分数'] = min(55+(patent_info['独立权利要求'].iloc[0].count("。"))*10,100)
        legal_status = patent_info['法律状态/事件']
        if legal_status.str.contains('授权', case=False, na=False).iloc[0]:
            legal_score = 85
        elif legal_status.str.contains('实质审查', case=False, na=False).iloc[0]:
            legal_score = 70
        temp_data_json['法律状态得分'] = legal_score
        patent_type= patent_info['专利类型']
        if patent_type.str.contains('发明', case=False, na=False).iloc[0]:
            patent_type_score = 90
        if patent_type.str.contains('实用新型', case=False, na=False).iloc[0]:
            patent_type_score = 50
        temp_data_json['专利类型得分'] = patent_type_score
        current_date = datetime.now()
        if not patent_info['预估到期日'].empty and patent_info['预估到期日'].iloc[0]!='-':
            year = pd.to_datetime(patent_info['预估到期日']).apply(lambda x: x.year - current_date.year - ((x.month, x.day) < (current_date.month, current_date.day)) if pd.notnull(x) else 5).iloc[0]
            if year>=16:
                year_score = 100
            elif 14 <= year <= 15:
                year_score = 90
            elif 12 <= year <= 13:
                year_score = 80
            elif 10 <= year < 11:
                year_score = 70
            elif 8 <= year <= 9:
                year_score = 60
            elif 6 <= year <= 7:
                year_score = 50
            temp_data_json['年限分数'] = year_score
    # 计算说明书分数
    # 打开 PDF 文件
    file_model = FileModel.objects.filter(pk=zarm.file_id)
    des_score = 75
    if file_model:
        doc = pymupdf.open(FileModel.objects.get(pk=zarm.file_id).file_path)
        # # 获取大纲
        # 搜索匹配标题的页数
        # matching_pages = set()
        # for item in outline:
        #     level, title, page = item
        #     if "DES".lower() in title.lower():  # 忽略大小写比较
        #         matching_pages.add(page)
        # des_num = len(matching_pages)
        des_num = doc.page_count
        def_num = 13
        if des_num > def_num:
            des_score = des_score+(des_num - def_num)*1.5
        else:
            des_score = des_score-(def_num-des_num)*4
    temp_data_json['说明书分数'] = min(des_score,100)

def get_clazz_type(p_code)->dict:
    """
    获取国民经济分类
    :param p_code: 不带最后一位
    :return: {"A":{"big_name":"xxx","category_list":{xxx}}}
    """

    class_info = df[df['公开(公告)号'] == p_code + "B"]
    if class_info.empty:
        class_info = df[df['公开(公告)号'] == p_code + "A"]
    if class_info.empty:
        class_info = df[df['公开(公告)号'] == p_code + "C"]
    if class_info.empty:
        class_info = df[df['公开(公告)号'] == p_code + "U"]
    if class_info.empty:
        class_info = df[df['公开(公告)号'] == p_code + "S"]
    clazz_map = {}
    if not class_info.empty:
        class_list = [item for item in class_info['国民经济行业分类号'].iloc[0].split("|") if item]
        if not class_info.empty:
            for index, clazz in enumerate(class_list):
                clazz = clazz.strip()
                big = clazz[0].upper()
                category = clazz[0:3].upper()
                class_name_df_item = class_name_df[class_name_df.iloc[:, 0] == big]
                if not class_name_df_item.empty:
                    big_name = class_name_df_item.iloc[:, 1].iloc[0]
                    category_name = class_name_df[class_name_df.iloc[:, 0] == category].iloc[:, 1].iloc[0].strip()
                    if big not in clazz_map:
                        clazz_map[big] = {"big_name":big_name,"category_set":{category_name,}}
                    else:
                        clazz_map[big]['category_set'].add(category_name)
    for value in clazz_map.values():
        new_category_list = []
        for i,name in enumerate(value['category_set']):
            if len(value['category_set'])>1:
                new_category_list.append(f"（{i+1}）{name}")
            else:
                new_category_list.append(f"{name}")
        value["category_list"] = new_category_list
    return clazz_map